"""
GENERADOR DE PLIEGOS (PPT): ppt_generation.py
DESCRIPCIÓN:
Este módulo se encarga de redactar documentos técnicos (Pliegos de Prescripciones Técnicas)
basándose en contratos previos similares.

Funciona en variar fases:
1. PLANIFICACIÓN: El LLM decide si faltan datos (presupuesto, uso, características).
2. REFERENCIA: Busca en Neo4j un contrato similar que tenga PPT (Pliego) adjunto.
3. GENERACIÓN: Le pasa al LLM la estructura del contrato antiguo para que la use de esqueleto.
4. EXPORTACIÓN: Convierte el texto Markdown resultante a un archivo Word (.docx).
"""

import re
from typing import Any, Dict, List, Optional, Tuple

import config
from clients import llm_client
from services.embeddings import embed_text
from services.neo4j_queries import neo4j_query, search_capitulos, search_extractos
from chat_utils.json_utils import safe_json_loads
from chat_utils.text_utils import clip
from chat_utils.prompt_loader import load_prompt

# Intentamos importar librería python-docx para crear Word. Si falla, el bot funcionará pero sin exportar archivo.
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


def plan_ppt_clarifications(user_request: str) -> Dict[str, Any]:
    """
    Analiza si la petición del usuario ("Hazme un pliego para un coche") es suficiente
    o si faltan detalles importantes para escribir algo decente.
    """
    prompt = load_prompt(
        "ppt_clarification",
        today=config.TODAY_STR,
        user_request=user_request
    )
    resp = llm_client.chat.completions.create(
        model=config.LLM_MODEL,
        messages=[{"role": "system", "content": "Devuelve SOLO JSON."}, {"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=500,
    )
    data = safe_json_loads(resp.choices[0].message.content or "") or {}
    
    need = bool(data.get("need_clarification"))
    normalized = data.get("normalized_request") or user_request
    questions = data.get("questions") if isinstance(data.get("questions"), list) else []
    questions = questions[:7] # Limitamos preguntas para no aburrir
    
    return {
        "need_clarification": need,
        "normalized_request": normalized,
        "questions": questions,
    }


def find_reference_ppt_contract(question_embedding: List[float], top_k: int = 10) -> Optional[Dict[str, Any]]:
    """
    Busca en el grafo el contrato "más parecido" que tenga un documento PPT (tipo_doc='PPT').
    Usa búsqueda vectorial sobre los capítulos.
    """
    # 1. Buscamos capítulos que se parezcan a la idea del usuario
    candidatos = search_capitulos(question_embedding, k=top_k, doc_tipo="PPT")
    
    # 2. De los capítulos encontrados, verificamos cuál pertenece a un PPT válido en la base de datos
    for c in candidatos:
        cid = c.get("contract_id")
        if not cid:
            continue
            
        rows = neo4j_query(
            """
            MATCH (c:ContratoRAG)-[td:TIENE_DOC]->(d:DocumentoRAG)
            WHERE (c.contract_id = $cid OR c.expediente = $cid) AND td.tipo_doc = 'PPT'
            RETURN d.doc_id AS doc_id
            LIMIT 1
            """,
            {"cid": cid},
        )
        if rows:
            c["doc_id"] = rows[0]["doc_id"]
            return c # Devolvemos el primer candidato válido
            
    return None


def get_ppt_reference_data(contract_id: str) -> Optional[Dict[str, Any]]:
    """
    Recupera TODOS los capítulos del PPT de referencia, ordenados.
    Esto sirve para que el LLM copie la estructura (Índice, apartados legales, técnicos...).
    """
    rows = neo4j_query(
        """
        MATCH (c:ContratoRAG)-[td:TIENE_DOC]->(d:DocumentoRAG)
        WHERE (c.contract_id = $cid OR c.expediente = $cid) AND td.tipo_doc = 'PPT'
        OPTIONAL MATCH (d)-[:TIENE_CAPITULO]->(cap:Capitulo)
        RETURN
          c.titulo     AS contrato_titulo,
          c.expediente AS expediente,
          c.contract_uri AS link_contrato,
          d.doc_id     AS doc_id,
          cap.heading  AS heading,
          cap.orden    AS orden,
          cap.texto    AS texto
        ORDER BY cap.orden ASC
        """,
        {"cid": contract_id},
    )
    if not rows:
        return None

    cap_list = []
    for r in rows:
        if r.get("heading") is None:
            continue
        cap_list.append({
            "heading": r.get("heading"),
            "orden": r.get("orden"),
            "texto": r.get("texto") or "",
            # NOTA: En el futuro podríamos filtrar textos muy largos aquí
        })

    return {
        "contract_id": contract_id,
        "expediente": rows[0].get("expediente"),
        "contrato_titulo": rows[0].get("contrato_titulo"),
        "link_contrato": rows[0].get("link_contrato"),
        "doc_id": rows[0].get("doc_id"),
        "capitulos": cap_list,
    }


def build_ppt_generation_prompt_one_by_one(user_request: str, ref_data: Dict[str, Any]) -> Tuple[str, str]:
    """
    Construye el 'Megaprompt' para que el LLM escriba el documento.
    Le inyectamos los capítulos de referencia para que los use de "inspiración estructural".
    """
    exp = ref_data.get("expediente") or "N/D"
    titulo_ref = ref_data.get("contrato_titulo") or "N/D"
    caps = ref_data.get("capitulos") or []

    cap_blocks = []
    for c in caps[:8]: # Usamos 8 capítulos para tener buena estructura
        heading = c.get("heading") or "N/D"
        orden = c.get("orden")
        texto = c.get("texto") or ""
        snippet = clip(texto, 200) # Solo tomamos el inicio para dar contexto de qué trata
        cap_blocks.append(
            f"### {orden}. {heading}\n"
            f"(Inicio: {snippet}...)"
        )
    caps_ref_text = "\n".join(cap_blocks) if cap_blocks else "N/D"

    system_msg = load_prompt("ppt_generation_system")
    system_msg += "\n\nIMPORTANTE: Tu respuesta DEBE ser ÚNICAMENTE el contenido del documento en formato Markdown. DEBE comenzar con '# Título del Documento'. Termina el documento de forma clara. Si sientes que te repites, DETENTE. NO escribas 'indefinidamente' ni entres en bucles."

    user_msg = load_prompt(
        "ppt_generation_user",
        today=config.TODAY_STR,
        user_request=user_request,
        exp=exp,
        judul_ref=titulo_ref, # Typo fix if needed, but keeping orig var name
        titulo_ref=titulo_ref,
        caps_ref_text=caps_ref_text
    )
    return system_msg.strip(), user_msg.strip()


def slug_filename(title: str, max_len: int = 80) -> str:
    """Convierte un título de documento ("Hola Mundo") en un nombre de archivo seguro ("hola-mundo")."""
    t = (title or "PPT").strip().lower()
    t = re.sub(r"[^\w\s-]", "", t, flags=re.UNICODE)
    t = re.sub(r"\s+", "-", t).strip("-")
    if len(t) > max_len:
        t = t[:max_len].rstrip("-")
    return t or "ppt-generado"


def ppt_to_docx_bytes(md_text: str, title: str = "Pliego de Prescripciones Técnicas") -> bytes:
    """
    Convierte el texto Markdown generado por el LLM a un archivo binario .docx (Word).
    Interpreta encabezados (##) y listas básicas.
    """
    if not HAS_DOCX:
        return b""
        
    doc = Document()
    doc.add_heading(title, level=0)
    
    # Parseo simple línea a línea
    for line in (md_text or "").splitlines():
        line = line.rstrip()
        if not line:
            continue
        # Título principal ya puesto, ignoramos si se repite al inicio
        if line.startswith("# ") and title.lower() in line.lower():
            continue
            
        # Subtítulos
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=2)
        elif line.startswith("- ") or line.startswith("* "):
            # Listas
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line) # Párrafo normal
            
    # Guardar en memoria (BytesIO) para enviarlo sin escribir en disco
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

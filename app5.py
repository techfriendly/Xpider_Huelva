# app.py
import os
import re
import json
import textwrap
from datetime import date
from typing import List, Dict, Any, Optional, Tuple
from collections import defaultdict

import chainlit as cl
from neo4j import GraphDatabase
from openai import OpenAI

# -----------------------------------------------------------------------------
# Opcional: exportar a Word
# -----------------------------------------------------------------------------
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# =============================================================================
# CONFIG
# =============================================================================

TODAY_STR = date.today().isoformat()

NEO4J_URI      = os.getenv("NEO4J_URI", "bolt://49.13.151.49:7687")
NEO4J_USER     = os.getenv("NEO4J_USER", "neo4j")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD", "root_tech_2019")
NEO4J_DB       = os.getenv("NEO4J_DB", "huelva")

LLM_BASE_URL = os.getenv("LLM_BASE_URL", "http://100.71.46.94:8002/v1")
LLM_API_KEY  = os.getenv("LLM_API_KEY", "dummy-key")
LLM_MODEL    = os.getenv("LLM_MODEL", "llm")

EMB_BASE_URL = os.getenv("EMB_BASE_URL", "http://100.71.46.94:8003/v1")
EMB_API_KEY  = os.getenv("EMB_API_KEY", "dummy-key")
EMB_MODEL    = os.getenv("EMB_MODEL", "embedding")
EMB_DIM      = int(os.getenv("EMB_DIM", "1024"))

K_CONTRATOS = int(os.getenv("K_CONTRATOS", "5"))
K_CAPITULOS = int(os.getenv("K_CAPITULOS", "10"))
K_EXTRACTOS = int(os.getenv("K_EXTRACTOS", "25"))

MAX_HISTORY_TURNS = int(os.getenv("MAX_HISTORY_TURNS", "6"))
SUGGESTION_LABEL_MAX_CHARS = int(os.getenv("SUGGESTION_LABEL_MAX_CHARS", "100"))

MODEL_MAX_CONTEXT_TOKENS = int(os.getenv("MODEL_MAX_CONTEXT_TOKENS", "12288"))
RESERVE_FOR_ANSWER_TOKENS = int(os.getenv("RESERVE_FOR_ANSWER_TOKENS", "1400"))
MEMORY_SUMMARY_TOKENS = int(os.getenv("MEMORY_SUMMARY_TOKENS", "250"))

RAG_CONTEXT_MAX_TOKENS = int(os.getenv("RAG_CONTEXT_MAX_TOKENS", "5500"))
RAG_CONTEXT_MAX_CHARS  = RAG_CONTEXT_MAX_TOKENS * 4

KNOWN_EXTRACTO_TYPES = [
    "normativas",
    "ubicaciones",
    "presupuesto_base",
    "garantia_definitiva",
    "garantia_otros_tipos",
    "duracion_y_prorrogas",
    "solvencia_tecnica",
    "solvencia_economica",
    "modificacion_contrato",
    "causas_imprevistas",
    "criterios_adjudicacion",
    "criterios_ambientales",
    "clausulas_sociales",
    "clausulas_igualdad_genero",
    "medios_personales",
    "medios_materiales",
    "CPV",
]


# =============================================================================
# CLIENTS
# =============================================================================

driver = GraphDatabase.driver(NEO4J_URI, auth=(NEO4J_USER, NEO4J_PASSWORD))
llm_client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
emb_client = OpenAI(base_url=EMB_BASE_URL, api_key=EMB_API_KEY)


# =============================================================================
# HELPERS: NEO4J
# =============================================================================

def neo4j_query(cypher: str, params: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    if params is None:
        params = {}
    with driver.session(database=NEO4J_DB) as session:
        res = session.run(cypher, **params)
        return [r.data() for r in res]


# =============================================================================
# HELPERS: JSON robusto
# =============================================================================

def _strip_code_fences(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"^```[a-zA-Z]*\s*", "", s)
    s = re.sub(r"\s*```$", "", s)
    s = s.replace("```", "").strip()
    i = s.find("{")
    j = s.rfind("}")
    if i >= 0 and j > i:
        s = s[i:j+1]
    return s.strip()

def safe_json_loads(s: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(_strip_code_fences(s))
    except Exception:
        return None


# =============================================================================
# HELPERS: texto / tokens / trimming
# =============================================================================

def clip(s: str, max_chars: int) -> str:
    s = s or ""
    return s if len(s) <= max_chars else s[:max_chars] + " […]"

def enforce_budget(text: str, max_chars: int) -> str:
    return text if len(text) <= max_chars else text[:max_chars] + " […]"

def estimate_tokens(text: str) -> int:
    if not text:
        return 0
    return max(1, len(text) // 4)

def trim_history_to_fit(
    history: List[Dict[str, str]],
    system_msg: str,
    user_msg: str,
    max_context_tokens: int = MODEL_MAX_CONTEXT_TOKENS,
    reserve_for_answer: int = RESERVE_FOR_ANSWER_TOKENS,
) -> List[Dict[str, str]]:
    budget = max_context_tokens - reserve_for_answer
    used = estimate_tokens(system_msg) + estimate_tokens(user_msg)
    trimmed: List[Dict[str, str]] = []
    for m in reversed(history):
        mt = estimate_tokens(m.get("content", ""))
        if used + mt > budget:
            break
        trimmed.append(m)
        used += mt
    return list(reversed(trimmed))

def context_token_report(system_msg: str, history: List[Dict[str, str]], user_msg: str) -> Dict[str, int]:
    sys_t = estimate_tokens(system_msg)
    hist_t = sum(estimate_tokens(m.get("content", "")) for m in history)
    user_t = estimate_tokens(user_msg)
    total = sys_t + hist_t + user_t
    return {"system": sys_t, "history": hist_t, "user": user_t, "total": total}


# =============================================================================
# HELPERS: embeddings
# =============================================================================

def embed_text(text: str, max_chars: int = 4000) -> List[float]:
    if not text:
        return []
    text = text[:max_chars]
    resp = emb_client.embeddings.create(model=EMB_MODEL, input=text)
    return resp.data[0].embedding


# =============================================================================
# VECTOR SEARCH EN NEO4J (RAG)
# =============================================================================

def search_contratos(embedding: List[float], k: int = K_CONTRATOS) -> List[Dict[str, Any]]:
    if not embedding:
        return []
    cypher = """
    CALL db.index.vector.queryNodes('contrato_rag_embedding', $k, $embedding)
    YIELD node, score
    OPTIONAL MATCH (e:EmpresaRAG)-[r:ADJUDICATARIA_RAG]->(node)
    RETURN
      node.contract_id                 AS contract_id,
      coalesce(node.expediente,'')     AS expediente,
      coalesce(node.titulo,'')         AS titulo,
      coalesce(node.abstract,'')       AS resumen,
      coalesce(node.estado,'')         AS estado,
      coalesce(node.cpv_principal,'')  AS cpv_principal,
      e.nif                            AS adjudicataria_nif,
      e.nombre                         AS adjudicataria_nombre,
      node.presupuesto_sin_iva         AS presupuesto_sin_iva,
      node.valor_estimado              AS valor_estimado,
      coalesce(r.importe_adjudicado, r.importe, node.importe_adjudicado) AS importe_adjudicado,
      score
    ORDER BY score DESC
    """
    return neo4j_query(cypher, {"k": k, "embedding": embedding})


def search_capitulos(embedding: List[float], k: int = K_CAPITULOS, doc_tipo: Optional[str] = None) -> List[Dict[str, Any]]:
    if not embedding:
        return []
    cypher = """
    CALL db.index.vector.queryNodes('capitulo_embedding', $k, $embedding)
    YIELD node, score
    MATCH (c:ContratoRAG)-[td:TIENE_DOC]->(d:DocumentoRAG)-[:TIENE_CAPITULO]->(node)
    WHERE ($doc_tipo IS NULL OR td.tipo_doc = $doc_tipo)
    RETURN
      node.cap_id                   AS cap_id,
      coalesce(node.heading,'')     AS heading,
      coalesce(node.texto,'')       AS texto,
      coalesce(node.fuente_doc,'')  AS fuente_doc,
      c.contract_id                 AS contract_id,
      coalesce(c.expediente,'')     AS expediente,
      coalesce(c.titulo,'')         AS contrato_titulo,
      score
    ORDER BY score DESC
    """
    return neo4j_query(cypher, {"k": k, "embedding": embedding, "doc_tipo": doc_tipo})


def search_extractos(
    embedding: List[float],
    k: int = K_EXTRACTOS,
    tipos: Optional[List[str]] = None,
    doc_tipo: Optional[str] = None
) -> List[Dict[str, Any]]:
    if not embedding:
        return []
    cypher = """
    CALL db.index.vector.queryNodes('extracto_embedding', $k, $embedding)
    YIELD node, score
    WITH node, score
    WHERE ($tipos IS NULL OR size($tipos)=0 OR node.tipo IN $tipos)
    MATCH (c:ContratoRAG)-[td:TIENE_DOC]->(d:DocumentoRAG)-[:TIENE_EXTRACTO]->(node)
    WHERE ($doc_tipo IS NULL OR td.tipo_doc = $doc_tipo)
    RETURN
      node.extracto_id               AS extracto_id,
      coalesce(node.tipo,'')         AS tipo,
      coalesce(node.texto,'')        AS texto,
      coalesce(node.fuente_doc,'')   AS fuente_doc,
      c.contract_id                  AS contract_id,
      coalesce(c.expediente,'')      AS expediente,
      coalesce(c.titulo,'')          AS contrato_titulo,
      score
    ORDER BY score DESC
    """
    return neo4j_query(cypher, {"k": k, "embedding": embedding, "tipos": tipos or [], "doc_tipo": doc_tipo})


# =============================================================================
# EVIDENCE MARKDOWN + SIDEBAR
# =============================================================================

def build_evidence_markdown(contratos, capitulos, extractos) -> str:
    lines: List[str] = []
    lines.append("### Evidencias utilizadas")

    if contratos:
        lines.append("")
        lines.append("**Contratos relevantes**")
        for c in contratos:
            lines.append(
                f"- Expediente **{c.get('expediente') or 'N/D'}** · "
                f"Título: {c.get('titulo') or 'N/D'} · "
                f"Adjudicataria: {c.get('adjudicataria_nombre') or 'N/D'} "
                f"(importe adjudicado: {c.get('importe_adjudicado') or 'N/D'})"
            )

    if capitulos:
        lines.append("")
        lines.append("**Capítulos relevantes**")
        for cap in capitulos:
            snippet = textwrap.shorten(cap.get("texto", "") or "", width=220, placeholder=" […]")
            lines.append(
                f"- Contrato **{cap.get('expediente') or 'N/D'}**, capítulo _{cap.get('heading') or 'N/D'}_ "
                f"({cap.get('fuente_doc') or ''}): {snippet}"
            )

    if extractos:
        lines.append("")
        lines.append("**Extractos relevantes**")
        for ex in extractos:
            snippet = textwrap.shorten(ex.get("texto", "") or "", width=220, placeholder=" […]")
            lines.append(
                f"- Contrato **{ex.get('expediente') or 'N/D'}**, tipo _{ex.get('tipo') or 'N/D'}_ "
                f"({ex.get('fuente_doc') or ''}): {snippet}"
            )

    return "\n".join(lines)

async def set_evidence_sidebar(title: str, markdown: str, props_extra: Optional[Dict[str, Any]] = None):
    """
    Abre/actualiza el sidebar derecho con un CustomElement EvidencePanel.
    Chainlit permite controlar el sidebar desde Python con ElementSidebar. 
    El CustomElement se implementa en public/elements/EvidencePanel.jsx. 
    """
    if props_extra is None:
        props_extra = {}

    props = {
        "title": title,
        "markdown": markdown,
        **props_extra,
    }

    try:
        el = cl.CustomElement(name="EvidencePanel", props=props)
        await cl.ElementSidebar.set_title(title)
        await cl.ElementSidebar.set_elements([el])
    except Exception:
        # Fallback: si tu versión de Chainlit no soporta ElementSidebar,
        # al menos adjuntamos el elemento al mensaje siguiente como "inline".
        # (No lo mandamos al chat para no ensuciar; pero evitamos crash.)
        pass

async def clear_evidence_sidebar():
    try:
        await cl.ElementSidebar.set_elements([])
    except Exception:
        pass


# =============================================================================
# (1) ROUTER DE INTENCIÓN (LLM)
# =============================================================================

def detect_intent(question: str) -> Dict[str, Any]:
    prompt = f"""
Fecha actual: {TODAY_STR}

Eres un enrutador de intención para un asistente de contratación pública con Neo4j GraphRAG.
Devuelve SOLO JSON válido:

{{
  "intent": "RAG_QA" | "CYPHER_QA" | "GENERATE_PPT",
  "doc_tipo": null | "PPT" | "PCAP",
  "extracto_tipos": null | ["normativas","solvencia_tecnica"],
  "needs_aggregation": true | false
}}

Reglas:
- Si pide contar/sumar/ranking/top/veces/estadísticas => CYPHER_QA.
- Si pide redactar/generar/elaborar un PPT => GENERATE_PPT.
- En otro caso => RAG_QA.
- Si menciona normativa/solvencia/garantías/criterios/ubicaciones/presupuesto/duración => extracto_tipos relevante.
- Si menciona explícitamente PPT o PCAP => doc_tipo.

Tipos extracto conocidos:
{KNOWN_EXTRACTO_TYPES}

Pregunta:
\"\"\"{question}\"\"\"
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": "Devuelve SOLO JSON válido."},
                  {"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=220,
    )
    data = safe_json_loads(resp.choices[0].message.content or "") or {}

    intent = data.get("intent", "RAG_QA")
    if intent not in ("RAG_QA", "CYPHER_QA", "GENERATE_PPT"):
        intent = "RAG_QA"

    doc_tipo = data.get("doc_tipo")
    if doc_tipo not in (None, "PPT", "PCAP"):
        doc_tipo = None

    tipos = data.get("extracto_tipos")
    if isinstance(tipos, list):
        tipos = [t for t in tipos if t in KNOWN_EXTRACTO_TYPES]
        tipos = tipos or None
    else:
        tipos = None

    needs_aggregation = bool(data.get("needs_aggregation", intent == "CYPHER_QA"))

    return {
        "intent": intent,
        "doc_tipo": doc_tipo,
        "extracto_tipos": tipos,
        "needs_aggregation": needs_aggregation,
    }


# =============================================================================
# (4) CYPHER QA (solo lectura) con autorepair “r”
# =============================================================================

WRITE_KEYWORDS = re.compile(
    r"\b(CREATE|MERGE|SET|DELETE|DETACH|DROP|LOAD\s+CSV|CALL\s+apoc\.|CALL\s+dbms)\b",
    re.IGNORECASE,
)

def cypher_is_safe_readonly(cypher: str) -> bool:
    if not cypher or not isinstance(cypher, str):
        return False
    if WRITE_KEYWORDS.search(cypher):
        return False
    if not re.search(r"\b(MATCH|CALL)\b", cypher, re.IGNORECASE):
        return False
    return True

def cypher_ensure_limit(cypher: str, default_limit: int = 50) -> str:
    if re.search(r"\bLIMIT\b", cypher, re.IGNORECASE):
        return cypher
    return cypher.rstrip() + f"\nLIMIT {default_limit}"

def cypher_needs_r_binding(cypher: str) -> bool:
    if re.search(r"\br\.\w+", cypher):
        return not bool(re.search(r"\[\s*r\s*:", cypher))
    return False

def get_schema_hint(max_chars: int = 7000) -> str:
    try:
        rows = neo4j_query("CALL db.schema.visualization()")
        if rows:
            return json.dumps(rows[0], ensure_ascii=False)[:max_chars]
    except Exception:
        pass
    return "N/D"

def generate_cypher_plan(question: str, schema_hint: str, error_hint: str = "") -> Dict[str, Any]:
    prompt = f"""
Fecha actual: {TODAY_STR}

Eres un experto en Neo4j y contratación pública.
Genera una consulta Cypher SOLO LECTURA para responder a la pregunta.

Esquema (puede estar truncado):
\"\"\"{schema_hint}\"\"\"

REGLAS IMPORTANTES:
- Si usas propiedades de la relación de adjudicación (r.importe_adjudicado o r.importe),
  DEBES declarar la relación con variable r, por ejemplo:
  (emp:EmpresaRAG)-[r:ADJUDICATARIA_RAG]->(c:ContratoRAG)

- Evita APOC.
- Evita CREATE/MERGE/SET/DELETE/DROP.

{("Error previo a corregir: " + error_hint) if error_hint else ""}

Devuelve SOLO JSON:
{{
  "cypher": "...",
  "params": {{}}
}}

Pregunta:
\"\"\"{question}\"\"\"
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": "Devuelve SOLO JSON válido."},
                  {"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=650,
    )
    data = safe_json_loads(resp.choices[0].message.content or "") or {}
    cypher = (data.get("cypher") or "").strip()
    params = data.get("params") or {}
    if not isinstance(params, dict):
        params = {}
    return {"cypher": cypher, "params": params, "raw": data}

def cypher_qa(question: str) -> Dict[str, Any]:
    schema_hint = get_schema_hint(7000)

    plan = generate_cypher_plan(question, schema_hint)
    cypher = plan["cypher"]
    params = plan["params"]

    if not cypher_is_safe_readonly(cypher):
        return {"error": "Cypher no seguro o inválido.", "cypher": cypher, "plan": plan}

    if cypher_needs_r_binding(cypher):
        plan = generate_cypher_plan(question, schema_hint, error_hint="La query usa r.<prop> pero no declara [r:REL].")
        cypher = plan["cypher"]
        params = plan["params"]
        if not cypher_is_safe_readonly(cypher):
            return {"error": "Cypher no seguro tras reparación.", "cypher": cypher, "plan": plan}

    cypher = cypher_ensure_limit(cypher, 50)

    try:
        rows = neo4j_query(cypher, params)
    except Exception as e:
        err = str(e)
        plan2 = generate_cypher_plan(question, schema_hint, error_hint=err)
        cypher2 = plan2["cypher"]
        params2 = plan2["params"]
        if not cypher_is_safe_readonly(cypher2):
            return {"error": f"Fallo Cypher y reparación insegura: {err}", "cypher": cypher, "plan": plan2}
        if cypher_needs_r_binding(cypher2):
            plan3 = generate_cypher_plan(question, schema_hint, error_hint="Define la relación con variable r si usas r.<prop>.")
            cypher2 = plan3["cypher"]
            params2 = plan3["params"]
            if not cypher_is_safe_readonly(cypher2):
                return {"error": "No se pudo generar Cypher seguro tras 2 reparaciones.", "cypher": cypher, "plan": plan3}
        cypher2 = cypher_ensure_limit(cypher2, 50)
        rows = neo4j_query(cypher2, params2)
        cypher = cypher2
        plan = plan2

    system_msg = (
        "Eres un asistente de contratación pública. Respondes SOLO con los datos devueltos por Neo4j. "
        "Si no hay datos suficientes, dilo."
    )
    user_msg = f"""
Pregunta:
\"\"\"{question}\"\"\"

Cypher ejecutado:
\"\"\"{cypher}\"\"\"

Filas devueltas (máx 50):
{json.dumps(rows[:50], ensure_ascii=False, indent=2)}

Responde en castellano, claro y conciso.
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
        temperature=0.2,
        max_tokens=700,
    )
    answer = (resp.choices[0].message.content or "").strip()

    return {"answer": answer, "cypher": cypher, "rows": rows, "plan": plan}


# =============================================================================
# Memoria corta
# =============================================================================

def summarize_for_memory(text: str, max_tokens: int = MEMORY_SUMMARY_TOKENS) -> str:
    if not text:
        return ""
    prompt = f"""
Resume el siguiente texto en un máximo de {max_tokens} tokens.
- Mantén datos clave (expediente, importes, requisitos, criterios) si aparecen.
- Si es un pliego, resume: objeto, alcance, estructura de capítulos y puntos críticos.
Devuelve SOLO el resumen.

TEXTO:
\"\"\"{text}\"\"\"
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[
            {"role": "system", "content": "Eres un asistente que resume de forma concisa."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
        max_tokens=max_tokens + 50,
    )
    return (resp.choices[0].message.content or "").strip()


# =============================================================================
# Follow-ups: solo si tiene sentido + robusto
# =============================================================================

def should_generate_followups(answer: str, contratos, capitulos, extractos) -> bool:
    if not answer or len(answer.strip()) < 200:
        return False
    if (len(contratos) + len(capitulos) + len(extractos)) == 0:
        return False
    return True

def generate_follow_up_questions(question: str, answer: str, max_suggestions: int = 3) -> List[str]:
    prompt = f"""
Genera entre 1 y {max_suggestions} preguntas de seguimiento útiles.
Devuelve EXCLUSIVAMENTE un JSON válido con la forma:

{{
  "suggestions": [
    "Pregunta 1",
    "Pregunta 2"
  ]
}}

Reglas:
- Cada suggestion debe ser un string JSON entre comillas.
- No uses ``` ni markdown.
- No repitas la pregunta original.

Pregunta:
\"\"\"{question}\"\"\"

Respuesta:
\"\"\"{answer}\"\"\"
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": "Devuelve SOLO JSON válido."},
                  {"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=256,
    )
    data = safe_json_loads(resp.choices[0].message.content or "")
    if not data or not isinstance(data.get("suggestions"), list):
        return []

    uniq: List[str] = []
    seen = set()
    for s in data["suggestions"]:
        if not isinstance(s, str):
            continue
        s_clean = " ".join(s.split()).strip()
        if not s_clean:
            continue
        if s_clean.lower() == question.strip().lower():
            continue
        if s_clean.lower() in seen:
            continue
        seen.add(s_clean.lower())
        uniq.append(s_clean)
        if len(uniq) >= max_suggestions:
            break
    return uniq


# =============================================================================
# PPT: clarificación + referencia + streaming + recomendaciones + evidencia + Word
# =============================================================================

def plan_ppt_clarifications(user_request: str) -> Dict[str, Any]:
    prompt = f"""
Fecha actual: {TODAY_STR}

Eres un analista de requisitos para redactar un PPT de contratación pública.
Devuelve SOLO JSON:

{{
  "need_clarification": true|false,
  "questions": ["..."],
  "normalized_request": "..."
}}

Criterio:
- Si faltan datos críticos (objeto exacto, alcance, localización, duración/plazos, presupuesto orientativo, SLAs/medios, entregables),
  pregunta (3 a 7 preguntas).
- Si está suficientemente claro, need_clarification=false.

Petición:
\"\"\"{user_request}\"\"\"
"""
    resp = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": "Devuelve SOLO JSON válido."},
                  {"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=500,
    )
    data = safe_json_loads(resp.choices[0].message.content or "") or {}
    need = bool(data.get("need_clarification", False))
    qs = data.get("questions", []) or []
    if not isinstance(qs, list):
        qs = []
    qs = [q for q in qs if isinstance(q, str) and q.strip()]
    nr = (data.get("normalized_request") or user_request).strip()
    return {"need_clarification": need, "questions": qs[:7], "normalized_request": nr}

def find_reference_ppt_contract(question_embedding: List[float], top_k: int = 10) -> Optional[Dict[str, Any]]:
    candidatos = search_contratos(question_embedding, top_k)
    if not candidatos:
        return None
    for c in candidatos:
        cid = c.get("contract_id")
        if not cid:
            continue
        rows = neo4j_query(
            """
            MATCH (c:ContratoRAG {contract_id: $cid})-[td:TIENE_DOC]->(d:DocumentoRAG)
            WHERE td.tipo_doc = 'PPT'
            RETURN d.doc_id AS doc_id
            LIMIT 1
            """,
            {"cid": cid}
        )
        if rows:
            c["doc_id"] = rows[0]["doc_id"]
            return c
    return None

def get_ppt_reference_data(contract_id: str) -> Optional[Dict[str, Any]]:
    rows = neo4j_query(
        """
        MATCH (c:ContratoRAG {contract_id: $cid})-[td:TIENE_DOC]->(d:DocumentoRAG)
        WHERE td.tipo_doc = 'PPT'
        OPTIONAL MATCH (d)-[:TIENE_CAPITULO]->(cap:Capitulo)
        RETURN
          c.titulo     AS contrato_titulo,
          c.expediente AS expediente,
          d.doc_id     AS doc_id,
          cap.heading  AS heading,
          cap.orden    AS orden,
          cap.texto    AS texto
        ORDER BY cap.orden ASC
        """,
        {"cid": contract_id}
    )
    if not rows:
        return None

    cap_list = []
    for r in rows:
        if r.get("heading") is None:
            continue
        cap_list.append({
            "heading": r.get("heading"),
            "orden": r.get("orden"),
            "texto": r.get("texto") or ""
        })

    return {
        "contract_id": contract_id,
        "expediente": rows[0].get("expediente"),
        "contrato_titulo": rows[0].get("contrato_titulo"),
        "doc_id": rows[0].get("doc_id"),
        "capitulos": cap_list,
    }

def build_ppt_generation_prompt_one_by_one(user_request: str, ref_data: Dict[str, Any]) -> Tuple[str, str]:
    exp = ref_data.get("expediente") or "N/D"
    titulo_ref = ref_data.get("contrato_titulo") or "N/D"
    caps = ref_data.get("capitulos") or []

    cap_blocks = []
    for c in caps[:18]:
        heading = c.get("heading") or "N/D"
        orden = c.get("orden")
        texto = c.get("texto") or ""
        snippet = clip(texto, 1200)
        cap_blocks.append(
            f"### Capítulo {orden}. {heading}\n"
            f"Contenido de referencia (no copiar literal):\n"
            f"{snippet}\n"
        )
    caps_ref_text = "\n".join(cap_blocks) if cap_blocks else "N/D"

    system_msg = (
        "Eres un redactor experto en Pliegos de Prescripciones Técnicas (PPT). "
        "Redactas de forma original, técnica y clara. "
        "Sigues la estructura del pliego de referencia y te inspiras en su contenido capítulo a capítulo, "
        "pero SIN copiar literal."
    )

    user_msg = f"""
Fecha actual: {TODAY_STR}

Se te pide redactar un **Pliego de Prescripciones Técnicas (PPT)**.

0) TITULACIÓN
- Comienza el documento con un título en H1:
  "# Pliego de Prescripciones Técnicas: <TÍTULO CONCRETO Y DESCRIPTIVO>"
  (El título debe reflejar el objeto real del encargo del usuario).

1) Encargo del usuario (objeto y contexto)
{user_request}

2) Pliego de referencia principal
- Expediente: {exp}
- Título del contrato de referencia: {titulo_ref}

3) Estructura y contenido del pliego de referencia (capítulo a capítulo)
Debes seguir la MISMA estructura de capítulos (mismos niveles), adaptando títulos y contenido al nuevo objeto.
NO copies literal. Para cada capítulo, redacta el capítulo equivalente para el nuevo PPT.

{caps_ref_text}

INSTRUCCIONES MUY IMPORTANTES:
- Redacta capítulo a capítulo en el mismo orden y con encabezados `##`.
- Tras cada capítulo, añade SIEMPRE un bloque en cursiva con el encabezado:
  _Recomendaciones para mejorar el pliego:_
  _- ..._
  _- ..._
  _- ..._
- Las recomendaciones deben ser prácticas: qué faltaría para mejorar ese capítulo en una licitación real.
- No inventes normativa nueva; si mencionas normativa, hazlo prudente y coherente.
- Devuelve SOLO el PPT en Markdown (sin comentarios fuera del pliego).
"""
    return system_msg.strip(), user_msg.strip()

def slug_filename(title: str, max_len: int = 80) -> str:
    t = (title or "PPT").strip().lower()
    t = re.sub(r"[^\w\s-]", "", t, flags=re.UNICODE)
    t = re.sub(r"\s+", "-", t).strip("-")
    if len(t) > max_len:
        t = t[:max_len].rstrip("-")
    return t or "ppt-generado"

def ppt_to_docx_bytes(md_text: str, title: str = "Pliego de Prescripciones Técnicas") -> bytes:
    if not HAS_DOCX:
        return b""
    doc = Document()
    doc.add_heading(title, level=1)
    for line in (md_text or "").splitlines():
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
        else:
            doc.add_paragraph(line)
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

async def handle_generate_ppt(question: str):
    plan = await cl.make_async(plan_ppt_clarifications)(question)
    if plan["need_clarification"] and plan["questions"]:
        cl.user_session.set("ppt_pending", True)
        cl.user_session.set("ppt_request_base", plan["normalized_request"])
        cl.user_session.set("ppt_questions", plan["questions"])

        qtxt = "\n".join([f"{i+1}. {q}" for i, q in enumerate(plan["questions"])])
        await cl.Message(
            content=(
                "Antes de redactar el PPT necesito aclarar algunas cosas:\n\n"
                f"{qtxt}\n\n"
                "Respóndeme en un solo mensaje (puedes numerar tus respuestas)."
            )
        ).send()
        return

    await cl.Message(content="Generando PPT basado en un pliego de referencia del grafo…").send()

    emb = await cl.make_async(embed_text)(question)
    if not emb:
        await cl.Message("No he podido calcular el embedding de la petición.").send()
        return

    ref_contrato = await cl.make_async(find_reference_ppt_contract)(emb, top_k=10)
    if not ref_contrato:
        await cl.Message("No he encontrado un PPT de referencia adecuado.").send()
        return

    contract_id = ref_contrato["contract_id"]
    ref_data = await cl.make_async(get_ppt_reference_data)(contract_id)
    if ref_data is None:
        await cl.Message(f"El contrato {contract_id} no tiene PPT con capítulos en el grafo.").send()
        return

    # Evidencia PPT para el panel lateral
    extra_caps = await cl.make_async(search_capitulos)(emb, k=min(12, K_CAPITULOS), doc_tipo="PPT")
    extra_extractos = await cl.make_async(search_extractos)(emb, k=min(20, K_EXTRACTOS), tipos=None, doc_tipo="PPT")
    evidence_md = build_evidence_markdown(
        contratos=[{
            "expediente": ref_data.get("expediente"),
            "titulo": ref_data.get("contrato_titulo"),
            "adjudicataria_nombre": ref_contrato.get("adjudicataria_nombre"),
            "importe_adjudicado": ref_contrato.get("importe_adjudicado"),
        }],
        capitulos=[{
            "heading": c.get("heading"),
            "expediente": ref_data.get("expediente"),
            "fuente_doc": "PPT",
            "texto": c.get("texto", ""),
        } for c in ref_data.get("capitulos", [])[:12]],
        extractos=[{
            "tipo": ex.get("tipo"),
            "expediente": ex.get("expediente"),
            "fuente_doc": ex.get("fuente_doc"),
            "texto": ex.get("texto", ""),
        } for ex in extra_extractos[:12]]
    )

    await set_evidence_sidebar(
        title="Evidencias RAG usadas (PPT)",
        markdown=evidence_md,
        props_extra={
            "mode": "PPT",
            "filters": {"doc_tipo": "PPT"},
            "counts": {"contratos": 1, "capitulos": len(ref_data.get("capitulos", [])[:12]), "extractos": len(extra_extractos[:12])},
        }
    )

    system_msg, user_msg = build_ppt_generation_prompt_one_by_one(question, ref_data)

    # Streaming PPT
    msg = cl.Message(content="")
    await msg.send()

    stream = llm_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
        max_tokens=6000,
        temperature=0.3,
        stream=True,
    )

    pliego_chunks: List[str] = []
    for chunk in stream:
        token = chunk.choices[0].delta.content or ""
        if token:
            pliego_chunks.append(token)
            await msg.stream_token(token)

    await msg.update()
    pliego_text = "".join(pliego_chunks).strip()

    # Título (H1)
    ppt_title = "Pliego de Prescripciones Técnicas"
    m = re.search(r"^#\s*(.+)$", pliego_text, flags=re.MULTILINE)
    if m:
        ppt_title = m.group(1).strip()

    # Word
    if HAS_DOCX:
        docx_bytes = ppt_to_docx_bytes(pliego_text, title=ppt_title)
        file = cl.File(
            name=f"{slug_filename(ppt_title)}.docx",
            content=docx_bytes,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        await cl.Message(content=f"Documento Word generado: **{ppt_title}**", elements=[file]).send()
    else:
        await cl.Message(content="(Aviso) No puedo generar Word porque python-docx no está disponible.").send()

    # Resumen corto
    ppt_summary = await cl.make_async(summarize_for_memory)(pliego_text, MEMORY_SUMMARY_TOKENS)
    await cl.Message(content=f"Resumen del PPT (memoria corta):\n\n{ppt_summary}").send()


# =============================================================================
# RAG CONTEXT BUILDER
# =============================================================================

def build_context(question: str, contratos, capitulos, extractos) -> str:
    parts: List[str] = []
    parts.append("=== PREGUNTA DEL USUARIO ===")
    parts.append(question.strip())

    if contratos:
        parts.append("\n=== CONTRATOS RELEVANTES ===")
        for c in contratos:
            snippet = clip(c.get("abstract", "") or "", 600)
            parts.append(
                f"- Expediente: {c.get('expediente') or 'N/D'} | Estado: {c.get('estado') or 'N/D'}\n"
                f"  Título: {c.get('titulo') or 'N/D'}\n"
                f"  CPV principal: {c.get('cpv_principal') or 'N/D'}\n"
                f"  Adjudicataria: {c.get('adjudicataria_nombre') or 'N/D'} "
                f"(NIF: {c.get('adjudicataria_nif') or 'N/D'})\n"
                f"  Presupuesto s/IVA: {c.get('presupuesto_sin_iva') or 'N/D'} | "
                f"Importe adjudicado: {c.get('importe_adjudicado') or 'N/D'}\n"
                f"  Resumen: {snippet}"
            )

    if capitulos:
        parts.append("\n=== CAPÍTULOS RELEVANTES ===")
        for cap in capitulos:
            snippet = clip(cap.get("texto", "") or "", 900)
            parts.append(
                f"- Contrato {cap.get('expediente') or 'N/D'} | Capítulo {cap.get('heading') or 'N/D'} "
                f"({cap.get('fuente_doc') or ''})\n"
                f"  Texto: {snippet}"
            )

    if extractos:
        parts.append("\n=== EXTRACTOS RELEVANTES ===")
        for ex in extractos:
            snippet = clip(ex.get("texto", "") or "", 700)
            parts.append(
                f"- Contrato {ex.get('expediente') or 'N/D'} | Tipo: {ex.get('tipo') or 'N/D'} "
                f"({ex.get('fuente_doc') or ''})\n"
                f"  Texto: {snippet}"
            )

    parts.append(
        "\n=== INSTRUCCIONES PARA EL MODELO ===\n"
        "Responde basándote EXCLUSIVAMENTE en el contexto anterior.\n"
        "Si no hay información suficiente, dilo.\n"
        "No inventes datos.\n"
        "Respuesta en castellano, clara y concisa."
    )

    ctx = "\n".join(parts)
    return enforce_budget(ctx, RAG_CONTEXT_MAX_CHARS)


# =============================================================================
# CHAINLIT
# =============================================================================

@cl.on_chat_start
async def on_chat_start():
    cl.user_session.set("history", [])
    cl.user_session.set("ppt_pending", False)
    cl.user_session.set("ppt_request_base", "")
    cl.user_session.set("ppt_questions", [])
    await clear_evidence_sidebar()

    await cl.Message(
        content=(
            "Hola. Soy el asistente RAG/Cypher de contratos (Huelva).\n\n"
            "Puedo:\n"
            "- Responder preguntas (RAG) y mostrar evidencias a la derecha.\n"
            "- Contar/sumar/rankings (Cypher).\n"
            "- Generar un PPT (te preguntaré si falta contexto) y descargarlo en Word.\n"
        )
    ).send()

@cl.action_callback("follow_up_question")
async def on_follow_up_question(action: cl.Action):
    payload = action.payload or {}
    q = payload.get("question")
    if not q:
        return
    await cl.Message(content=q).send()
    await on_message(cl.Message(content=q))

@cl.on_message
async def on_message(message: cl.Message):
    question = (message.content or "").strip()
    if not question:
        await cl.Message(content="No he recibido ninguna pregunta.").send()
        return

    # Si estamos en clarificación PPT, este mensaje son las respuestas
    if cl.user_session.get("ppt_pending", False):
        base_req = cl.user_session.get("ppt_request_base", "")
        final_req = f"{base_req}\n\nAclaraciones del usuario:\n{question}"
        cl.user_session.set("ppt_pending", False)
        cl.user_session.set("ppt_request_base", "")
        cl.user_session.set("ppt_questions", [])
        await handle_generate_ppt(final_req)
        return

    history: List[Dict[str, str]] = cl.user_session.get("history", [])
    thinking_msg = await cl.Message(content="Detectando intención y consultando el grafo...").send()

    try:
        intent = await cl.make_async(detect_intent)(question)

        # GENERATE_PPT
        if intent["intent"] == "GENERATE_PPT":
            await thinking_msg.update()
            await handle_generate_ppt(question)
            history.append({"role": "user", "content": question})
            history.append({"role": "assistant", "content": f"PPT generado (fecha {TODAY_STR}). Word entregado + resumen corto."})
            cl.user_session.set("history", history[-MAX_HISTORY_TURNS:])
            return

        # CYPHER_QA
        if intent["intent"] == "CYPHER_QA":
            thinking_msg.content = "Generando y ejecutando consulta Cypher (solo lectura)..."
            await thinking_msg.update()

            out = await cl.make_async(cypher_qa)(question)
            if out.get("error"):
                await cl.Message(content=f"No he podido ejecutar Cypher QA.\nDetalle: {out.get('error')}").send()
                return

            answer = out["answer"]
            await cl.Message(content=answer).send()

            history.append({"role": "user", "content": question})
            answer_mem = await cl.make_async(summarize_for_memory)(answer, MEMORY_SUMMARY_TOKENS) if len(answer) > 2000 else answer
            history.append({"role": "assistant", "content": answer_mem})
            cl.user_session.set("history", history[-MAX_HISTORY_TURNS:])

            thinking_msg.content = f"Respuesta generada (Cypher). Tokens aprox: enviados={estimate_tokens(question)}, generados={estimate_tokens(answer)}"
            await thinking_msg.update()
            return

        # RAG_QA
        thinking_msg.content = "Ejecutando RAG (vector search) con filtros..."
        await thinking_msg.update()

        embedding = await cl.make_async(embed_text)(question)
        if not embedding:
            thinking_msg.content = "No he podido generar el embedding."
            await thinking_msg.update()
            return

        doc_tipo = intent.get("doc_tipo")
        tipos = intent.get("extracto_tipos")

        contratos = await cl.make_async(search_contratos)(embedding, K_CONTRATOS)
        capitulos = await cl.make_async(search_capitulos)(embedding, K_CAPITULOS, doc_tipo)
        extractos = await cl.make_async(search_extractos)(embedding, K_EXTRACTOS, tipos, doc_tipo)

        evidence_md = build_evidence_markdown(contratos, capitulos, extractos)
        context = build_context(question, contratos, capitulos, extractos)

        system_msg = (
            "Eres un asistente experto en contratación pública. "
            "Respondes SOLO con la información del contexto. No inventas datos."
        )

        # recorta history a lo que quepa
        history_short = history[-MAX_HISTORY_TURNS:]
        history_trimmed = trim_history_to_fit(
            history=history_short,
            system_msg=system_msg,
            user_msg=context,
            max_context_tokens=MODEL_MAX_CONTEXT_TOKENS,
            reserve_for_answer=RESERVE_FOR_ANSWER_TOKENS
        )

        rep = context_token_report(system_msg, history_trimmed, context)

        thinking_msg.content = (
            f"Redactando respuesta… Tokens aprox enviados={rep['total']} "
            f"(sys={rep['system']}, hist={rep['history']}, ctx={rep['user']}). "
            f"Filtros: doc_tipo={doc_tipo}, extracto_tipos={tipos}"
        )
        await thinking_msg.update()

        # Sidebar: evidencias RAG (derecha)
        await set_evidence_sidebar(
            title="Evidencias RAG usadas",
            markdown=evidence_md,
            props_extra={
                "mode": "RAG",
                "filters": {"doc_tipo": doc_tipo, "extracto_tipos": tipos},
                "tokens": {"sent_approx": rep["total"], "budget": MODEL_MAX_CONTEXT_TOKENS},
                "counts": {"contratos": len(contratos), "capitulos": len(capitulos), "extractos": len(extractos)},
            }
        )

        messages_llm: List[Dict[str, str]] = [{"role": "system", "content": system_msg}]
        messages_llm.extend(history_trimmed)
        messages_llm.append({"role": "user", "content": context})

        reply = cl.Message(content="")
        await reply.send()

        stream = llm_client.chat.completions.create(
            model=LLM_MODEL,
            messages=messages_llm,
            temperature=0.3,
            stream=True,
            max_tokens=900,
        )

        full_answer: List[str] = []
        for chunk in stream:
            token = chunk.choices[0].delta.content or ""
            if token:
                full_answer.append(token)
                await reply.stream_token(token)

        await reply.update()
        answer = "".join(full_answer).strip()

        # memoria corta
        history.append({"role": "user", "content": question})
        answer_mem = await cl.make_async(summarize_for_memory)(answer, MEMORY_SUMMARY_TOKENS) if len(answer) > 2000 else answer
        history.append({"role": "assistant", "content": answer_mem})
        cl.user_session.set("history", history[-MAX_HISTORY_TURNS:])

        # follow-ups
        suggestions: List[str] = []
        if should_generate_followups(answer, contratos, capitulos, extractos):
            suggestions = await cl.make_async(generate_follow_up_questions)(question, answer, 3)

        actions: List[cl.Action] = []
        for s in suggestions:
            label = s if len(s) <= SUGGESTION_LABEL_MAX_CHARS else s[:SUGGESTION_LABEL_MAX_CHARS - 1] + "…"
            actions.append(cl.Action(
                name="follow_up_question",
                label=label,
                tooltip=s,
                payload={"question": s},
                icon="sparkles",
            ))

        reply.actions = actions
        await reply.update()

        gen_tokens = estimate_tokens(answer)
        thinking_msg.content = f"Respuesta generada (RAG). Tokens aprox: enviados={rep['total']}, generados={gen_tokens}"
        await thinking_msg.update()

    except Exception as e:
        print(f"[ERROR] {e}")
        thinking_msg.content = "Ha ocurrido un error. Revisa logs."
        await thinking_msg.update()
